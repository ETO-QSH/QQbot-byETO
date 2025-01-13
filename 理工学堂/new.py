import os
import json
import requests
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import RGBColor

base_url = "http://lgxt.wutp.com.cn/api"
headers = {
    'Accept': '*/*',
    'Accept-Language': 'zh-CN,zh;q=0.9',
    'Content-Type': 'application/x-www-form-urlencoded',
}

session = requests.Session()
InformationDictionary = {}


def login(username, password):
    login_url = f"{base_url}/login"
    data = {
        'loginName': username,
        'password': password,
    }
    try:
        response = session.post(login_url, headers=headers, data=data, timeout=10)
        response.raise_for_status()
        result = response.json()
        if result['code'] == 0:
            token = result['data']
            headers['Authorization'] = token
            session.headers.update({'Authorization': token})
            return True, "欢迎回来！"
        else:
            return False, f"登录失败：{result['msg']}"
    except requests.exceptions.RequestException as e:
        return False, f"网络错误：{e}"


def get_user_info():
    user_info_url = f"{base_url}/userInfo"
    try:
        response = session.post(user_info_url, headers=headers, timeout=10)
        response.raise_for_status()
        result = response.json()
        if result['code'] == 0:
            user_info = result['data']
            return True, user_info
        else:
            return False, f"获取用户信息失败：{result['msg']}"
    except requests.exceptions.RequestException as e:
        return False, f"网络错误：{e}"


def get_my_courses():
    my_courses_url = f"{base_url}/myCourses"
    try:
        response = session.post(my_courses_url, headers=headers, timeout=10)
        response.raise_for_status()
        result = response.json()
        if result['code'] == 0:
            courses = result['data']
            return True, courses
        else:
            return False, f"获取课程列表失败：{result['msg']}"
    except requests.exceptions.RequestException as e:
        return False, f"网络错误：{e}"


def get_course_works(course_id):
    my_course_works_url = f"{base_url}/myCourseWorks"
    data = {
        'courseId': course_id,
    }
    try:
        response = session.post(my_course_works_url, headers=headers, data=data, timeout=10)
        response.raise_for_status()
        result = response.json()
        if result['code'] == 0:
            works = result['data']
            return True, works
        else:
            return False, f"获取课程作业失败：{result['msg']}"
    except requests.exceptions.RequestException as e:
        return False, f"网络错误：{e}"


def get_questions(work_id):
    show_questions_url = f"{base_url}/showQuestions"
    data = {
        'workId': work_id,
    }
    try:
        response = session.post(show_questions_url, headers=headers, data=data, timeout=10)
        response.raise_for_status()
        result = response.json()
        if result['code'] == 0:
            questions = result['data']
            return True, questions
        else:
            return False, f"获取题目失败：{result['msg']}"
    except requests.exceptions.RequestException as e:
        return False, f"网络错误：{e}"


def submit_answer(work_id, grade):
    submit_answer_url = f"{base_url}/submitAnswer"
    data = {
        'grade': grade,
        'workId': work_id,
    }
    try:
        response = session.post(submit_answer_url, headers=headers, data=data, timeout=10)
        response.raise_for_status()
        result = response.json()
        if result['code'] == 0:
            return True, f"答案提交成功，成绩：{grade}\n返回信息：{result['data']}"
        else:
            return False, f"答案提交失败：{result['msg']}"
    except requests.exceptions.RequestException as e:
        return False, f"网络错误：{e}"


def collect_all_questions(work_id, max_iterations=100):
    collected_questions = {}
    no_new_questions_count = 0

    for i in range(max_iterations):
        success, questions = get_questions(work_id)
        if success:
            new_question_found = False
            for question in questions:
                question_id = question.get('id')
                if question_id not in collected_questions:
                    collected_questions[question_id] = question
                    new_question_found = True
            if new_question_found:
                no_new_questions_count = 0
            else:
                no_new_questions_count += 1
                if no_new_questions_count >= 10:
                    break
        else:
            break

    return collected_questions


def save_questions_to_word(collected_questions, assignment_folder, work_name, export_answers=True):
        document = Document()

        InformationDictionary[work_name] = {}

        style = document.styles['Normal']
        font_style = style.font
        font_style.name = '宋体'
        font_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        document.add_heading(work_name, 0)

        images_folder = os.path.join(assignment_folder, '题目图片')

        sorted_question_ids = sorted(collected_questions.keys(), key=lambda x: int(x))

        for idx, question_id in enumerate(sorted_question_ids):
            question = collected_questions[question_id]
            name = question.get('name', 'N/A')
            answer = question.get('answer', 'N/A')

            document.add_heading(f'题目 {idx + 1}: {name}', level=2)

            InformationDictionary[work_name][question_id] = [name, answer]

            image_path = os.path.join(images_folder, f"{question_id}.png")
            if os.path.exists(image_path):
                document.add_picture(image_path, width=Inches(5))
            else:
                document.add_paragraph("（无图片）")

            if export_answers:
                answer_paragraph = document.add_paragraph("答案：")
                answer_run = answer_paragraph.add_run(answer)
                answer_run.font.color.rgb = RGBColor(255, 0, 0)

        docx_path = os.path.join(assignment_folder, f'{work_name}.docx')
        document.save(docx_path)


login_state, login_result = login('2373204754', '******')
info_state, info_result = get_user_info()

courses = {}
courses_state, courses_result = get_my_courses()
for item in courses_result:
    courses[item['bookName']] = item['courseId']
print(courses)

works = {}
works_state, works_result = get_course_works('42')
for item in works_result:
    works[item['workName']] = item['workId'], item['times'], item['grade']
print(works)

with open('高等数学.json', 'w', encoding='utf-8') as file:
    json.dump(InformationDictionary, file, indent=4)
